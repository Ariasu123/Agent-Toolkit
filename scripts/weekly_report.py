#!/usr/bin/env python3
"""通用周报生成工具。

这个脚本由 weekly-report skill 调用，负责：
- 扫描既有周报并维护 used.md
- 从 Crossref 中选择未重复的相近方向文献
- 解析用户给出的本周内容 Markdown
- 复制最新周报模板并写入周次、正文、图片和文献精读表
"""

from __future__ import annotations

import argparse
import datetime as dt
import html
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import urllib.parse
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from docx.text.paragraph import Paragraph


DEFAULT_REPORT_DIR = Path(os.environ.get("WEEKLY_REPORT_DIR", ".")).expanduser().resolve()
DEFAULT_REPORT_PREFIX = os.environ.get("WEEKLY_REPORT_PREFIX", "weekly-report")
DEFAULT_INPUT_NAME = "weekly-input.md"
DEFAULT_RENDER_GLOB = (
    Path.home()
    / ".codex/plugins/cache/openai-primary-runtime/documents/*/skills/documents/render_docx.py"
)

DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+\b", re.I)
WEEK_RE = re.compile(
    r"周次[:：]\s*(\d{4})\.(\d{1,2})\.(\d{1,2})\s*[-–]\s*(?:(\d{4})\.)?(\d{1,2})\.(\d{1,2})"
)
GENERIC_REPORT_NAME_RE = re.compile(r"(.+)-(\d{2})-?(\d{4})(?:-v(\d+))?\.docx$")

BODY_FONT = "Times New Roman"
EAST_ASIA_FONT = os.environ.get("WEEKLY_REPORT_CJK_FONT", "宋体")
BODY_SIZE = Pt(12)
HEADER_SIZE = Pt(14)


@dataclass
class PaperRecord:
    key: str
    doi: str = ""
    title: str = ""
    year: str = ""
    venue: str = ""
    citation: str = ""
    occurrences: list[str] = field(default_factory=list)


def report_dir_from_args(value: str | None) -> Path:
    return Path(value).expanduser().resolve() if value else DEFAULT_REPORT_DIR


def used_file_for(report_dir: Path, value: str | None = None) -> Path:
    env = os.environ.get("USED_FILE")
    if value:
        return Path(value).expanduser()
    if env:
        return Path(env).expanduser()
    return report_dir / "used.md"


def default_input_path(report_dir: Path) -> Path:
    return report_dir / DEFAULT_INPUT_NAME


def input_path_from_args(args, report_dir: Path) -> Path:
    if getattr(args, "input", None):
        path = Path(args.input).expanduser()
        if not path.is_absolute():
            path = Path.cwd() / path
        path = path.resolve()
        if not path.exists():
            raise SystemExit(f"输入文件不存在：{path}")
        return path

    path = default_input_path(report_dir)
    if not path.exists():
        raise SystemExit(
            "默认输入文件不存在："
            f"{path}\n"
            "请创建 weekly-input.md，或用 --input /path/to/weekly-input.md 指定其他输入文件。"
        )
    return path


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(text or "").replace("\u00a0", " ")).strip()


def report_prefix_from_args(value: str | None) -> str:
    prefix = normalize_space(value or os.environ.get("WEEKLY_REPORT_PREFIX", DEFAULT_REPORT_PREFIX))
    if not prefix:
        raise SystemExit("周报文件名前缀不能为空。")
    if "/" in prefix or "\\" in prefix:
        raise SystemExit("周报文件名前缀不能包含路径分隔符。")
    return prefix


def normalize_doi(doi: str) -> str:
    doi = normalize_space(doi)
    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi, flags=re.I)
    doi = doi.rstrip(".,;，。；)")
    return doi.lower()


def extract_doi(text: str) -> str:
    match = DOI_RE.search(text or "")
    return normalize_doi(match.group(0)) if match else ""


def normalize_title(text: str) -> str:
    text = normalize_space(text).lower()
    text = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", text)
    return text


def markdown_escape(text: str) -> str:
    text = normalize_space(text)
    return text.replace("|", "\\|")


def report_name_re(report_prefix: str) -> re.Pattern[str]:
    return re.compile(rf"{re.escape(report_prefix)}-(\d{{2}})-?(\d{{4}})(?:-v(\d+))?\.docx$")


def parse_report_date(path: Path, report_prefix: str | None = None) -> tuple[dt.date, int] | None:
    if report_prefix:
        match = report_name_re(report_prefix).match(path.name)
        if not match:
            return None
        yy, mmdd, version = match.groups()
    else:
        match = GENERIC_REPORT_NAME_RE.match(path.name)
        if not match:
            return None
        _prefix, yy, mmdd, version = match.groups()
    year = 2000 + int(yy)
    month = int(mmdd[:2])
    day = int(mmdd[2:])
    try:
        date = dt.date(year, month, day)
    except ValueError:
        return None
    return date, int(version or "1")


def report_files(report_dir: Path, report_prefix: str) -> list[Path]:
    files: list[tuple[dt.date, int, float, Path]] = []
    for path in report_dir.glob("*.docx"):
        parsed = parse_report_date(path, report_prefix)
        if not parsed:
            continue
        date, version = parsed
        files.append((date, version, path.stat().st_mtime, path))
    return [item[-1] for item in sorted(files)]


def latest_report(report_dir: Path, report_prefix: str) -> Path:
    files = report_files(report_dir, report_prefix)
    if not files:
        raise SystemExit(f"未找到匹配前缀 {report_prefix!r} 的周报 docx：{report_dir}")
    return files[-1]


def crossref_lookup(doi: str, timeout: float = 8.0) -> dict[str, Any] | None:
    doi = normalize_doi(doi)
    if not doi:
        return None
    url = "https://api.crossref.org/works/" + urllib.parse.quote(doi, safe="")
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "weekly-report-skill/0.1"})
        with urllib.request.urlopen(req, timeout=timeout) as response:
            data = json.loads(response.read().decode("utf-8"))
        return data.get("message")
    except Exception:
        return None


def crossref_search(query: str, from_year: int = 2021, rows: int = 12) -> list[dict[str, Any]]:
    params = {
        "query.bibliographic": query,
        "filter": f"from-pub-date:{from_year}-01-01,type:journal-article",
        "sort": "relevance",
        "order": "desc",
        "rows": str(rows),
    }
    url = "https://api.crossref.org/works?" + urllib.parse.urlencode(params)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "weekly-report-skill/0.1"})
        with urllib.request.urlopen(req, timeout=10) as response:
            data = json.loads(response.read().decode("utf-8"))
        return data.get("message", {}).get("items", [])
    except Exception:
        return []


def first_date_year(message: dict[str, Any]) -> str:
    for key in ("published-print", "published-online", "published", "created"):
        parts = message.get(key, {}).get("date-parts")
        if parts and parts[0]:
            return str(parts[0][0])
    return ""


def author_text(message: dict[str, Any], limit: int = 5) -> str:
    authors = []
    for author in message.get("author", [])[:limit]:
        given = author.get("given", "")
        family = author.get("family", "")
        name = normalize_space(f"{given} {family}")
        if name:
            authors.append(name)
    if not authors:
        return ""
    if len(message.get("author", [])) > limit:
        return ", ".join(authors) + " et al."
    return ", ".join(authors)


def metadata_to_paper(message: dict[str, Any], query: str = "", score: float = 0) -> dict[str, Any]:
    title = normalize_space((message.get("title") or [""])[0])
    venue = normalize_space((message.get("container-title") or [""])[0])
    doi = normalize_doi(message.get("DOI", ""))
    year = first_date_year(message)
    authors = author_text(message)
    pages = normalize_space(message.get("page", ""))
    volume = normalize_space(message.get("volume", ""))
    issue = normalize_space(message.get("issue", ""))
    vol_issue = volume + (f"({issue})" if issue else "")
    journal_part = ", ".join(part for part in [venue, vol_issue, pages] if part)
    citation = normalize_space(
        f"{authors}. {title}. {journal_part} ({year}). DOI: {doi}."
    ).replace("..", ".")
    return {
        "title": title,
        "doi": doi,
        "year": year,
        "venue": venue,
        "authors": authors,
        "url": message.get("URL") or (f"https://doi.org/{doi}" if doi else ""),
        "citation": citation,
        "query": query,
        "score": round(score, 3),
        "source": "Crossref",
        "abstract": normalize_space(re.sub("<[^>]+>", "", message.get("abstract", ""))),
    }


def parse_title_from_citation(citation: str) -> str:
    citation = normalize_space(citation)
    if not citation:
        return ""
    no_doi = re.sub(r"\bDOI[:：]?\s*" + DOI_RE.pattern, "", citation, flags=re.I)
    parts = [part.strip() for part in re.split(r"\.\s+", no_doi) if part.strip()]
    if len(parts) >= 2:
        return parts[1]
    return citation[:120]


def parse_year_from_citation(citation: str) -> str:
    years = re.findall(r"\b(20\d{2}|19\d{2})\b", citation or "")
    return years[-1] if years else ""


def extract_week_text(doc: Document) -> str:
    for para in doc.paragraphs:
        text = normalize_space(para.text)
        if text.startswith("周次"):
            return text
    return ""


def extract_citation(doc: Document) -> str:
    if not doc.tables or len(doc.tables[0].rows) < 2:
        return ""
    return normalize_space(doc.tables[0].rows[1].cells[0].text)


def occurrence_text(path: Path, week: str) -> str:
    week_clean = week.replace("周次：", "").replace("周次:", "").strip()
    return f"{week_clean} ({path.name})" if week_clean else path.name


def record_from_docx(path: Path) -> PaperRecord | None:
    try:
        doc = Document(path)
    except Exception:
        return None
    citation = extract_citation(doc)
    if not citation:
        return None
    doi = extract_doi(citation)
    metadata = crossref_lookup(doi) if doi else None
    title = ""
    year = ""
    venue = ""
    if metadata:
        title = normalize_space((metadata.get("title") or [""])[0])
        year = first_date_year(metadata)
        venue = normalize_space((metadata.get("container-title") or [""])[0])
    if not title:
        title = parse_title_from_citation(citation)
    if not year:
        year = parse_year_from_citation(citation)
    key = f"doi:{doi}" if doi else f"title:{normalize_title(title or citation)}"
    week = extract_week_text(doc)
    return PaperRecord(
        key=key,
        doi=doi,
        title=title,
        year=year,
        venue=venue,
        citation=citation,
        occurrences=[occurrence_text(path, week)],
    )


def scan_used_records(report_dir: Path, report_prefix: str) -> list[PaperRecord]:
    merged: dict[str, PaperRecord] = {}
    for path in report_files(report_dir, report_prefix):
        record = record_from_docx(path)
        if not record:
            continue
        existing = merged.get(record.key)
        if existing is None:
            merged[record.key] = record
            continue
        for occurrence in record.occurrences:
            if occurrence not in existing.occurrences:
                existing.occurrences.append(occurrence)
        if not existing.title and record.title:
            existing.title = record.title
        if not existing.venue and record.venue:
            existing.venue = record.venue
        if not existing.year and record.year:
            existing.year = record.year
    return sorted(
        merged.values(),
        key=lambda r: (r.year or "0000", r.title.lower(), r.doi),
        reverse=True,
    )


def write_used_md(records: list[PaperRecord], used_file: Path) -> None:
    lines = [
        "# 已用文献",
        "",
        "| Key | DOI | Year | Title | Venue | Citation | Occurrences |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for record in records:
        lines.append(
            "| "
            + " | ".join(
                [
                    markdown_escape(record.key),
                    markdown_escape(record.doi),
                    markdown_escape(record.year),
                    markdown_escape(record.title),
                    markdown_escape(record.venue),
                    markdown_escape(record.citation),
                    markdown_escape("; ".join(record.occurrences)),
                ]
            )
            + " |"
        )
    used_file.parent.mkdir(parents=True, exist_ok=True)
    used_file.write_text("\n".join(lines) + "\n", encoding="utf-8")


def sync_used(report_dir: Path, used_file: Path, report_prefix: str) -> list[PaperRecord]:
    records = scan_used_records(report_dir, report_prefix)
    write_used_md(records, used_file)
    return records


def read_used_keys(used_file: Path) -> tuple[set[str], set[str]]:
    doi_keys: set[str] = set()
    title_keys: set[str] = set()
    if not used_file.exists():
        return doi_keys, title_keys
    for line in used_file.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| doi:") and not line.startswith("| title:"):
            continue
        cells = [cell.strip().replace("\\|", "|") for cell in line.strip("|").split("|")]
        if not cells:
            continue
        key = cells[0]
        if key.startswith("doi:"):
            doi_keys.add(key[4:].lower())
        elif key.startswith("title:"):
            title_keys.add(key[6:])
        if len(cells) > 1 and cells[1]:
            doi_keys.add(normalize_doi(cells[1]))
        if len(cells) > 3 and cells[3]:
            title_keys.add(normalize_title(cells[3]))
    return doi_keys, title_keys


def paper_score(paper: dict[str, Any]) -> float:
    title = paper.get("title", "").lower()
    abstract = paper.get("abstract", "").lower()
    text = f"{title} {abstract}"
    score = 0.0
    for term, weight in [
        ("mid-infrared", 5),
        ("mid infrared", 5),
        ("infrared", 2),
        ("gas", 4),
        ("spectrometer", 4),
        ("spectroscopy", 3),
        ("computational", 4),
        ("spectral reconstruction", 5),
        ("reconstruction", 3),
        ("metasurface", 3),
        ("filter", 2),
        ("miniature", 2),
        ("microspectrometer", 4),
        ("machine learning", 2),
    ]:
        if term in text:
            score += weight
    try:
        year = int(paper.get("year") or 0)
    except ValueError:
        year = 0
    if year >= 2021:
        score += 3
    if year >= 2025:
        score += 1
    if paper.get("doi"):
        score += 1
    return score


def select_paper(report_dir: Path, used_file: Path, report_prefix: str) -> dict[str, Any]:
    if not used_file.exists():
        sync_used(report_dir, used_file, report_prefix)
    used_dois, used_titles = read_used_keys(used_file)
    queries = [
        "mid-infrared gas sensing computational spectrometer",
        "mid-infrared computational spectroscopy gas sensing",
        "miniature spectrometer spectral reconstruction filter array",
        "mid-infrared metasurface microspectrometer gas sensing",
        "infrared gas sensor spectral reconstruction machine learning",
        "filter array computational spectrometer spectral reconstruction",
    ]
    candidates: dict[str, dict[str, Any]] = {}
    for from_year in (2021, 2016):
        for query in queries:
            for item in crossref_search(query, from_year=from_year):
                paper = metadata_to_paper(item, query=query)
                doi = paper.get("doi", "")
                title_key = normalize_title(paper.get("title", ""))
                if doi and doi in used_dois:
                    continue
                if title_key and title_key in used_titles:
                    continue
                if not paper.get("title"):
                    continue
                score = paper_score(paper)
                if from_year < 2021:
                    score -= 2
                paper["score"] = round(score, 3)
                candidate_key = doi or title_key
                if candidate_key and score > candidates.get(candidate_key, {}).get("score", -999):
                    candidates[candidate_key] = paper
        strong = [p for p in candidates.values() if p.get("score", 0) >= 8]
        if strong:
            break
    if not candidates:
        raise SystemExit("没有找到可用的未重复文献。")
    ranked = sorted(candidates.values(), key=lambda p: p.get("score", 0), reverse=True)
    best = ranked[0]
    if best.get("score", 0) < 5:
        raise SystemExit("没有找到高置信度的未重复相关文献。")
    best["alternatives"] = ranked[1:4]
    return best


def parse_weekly_input(input_path: Path) -> dict[str, Any]:
    text = input_path.read_text(encoding="utf-8")
    heading_re = re.compile(r"^##\s*(本周目标和方案|本周进展和问题|下周研究计划|其他)\s*$", re.M)
    matches = list(heading_re.finditer(text))
    sections: dict[str, str] = {}
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        sections[match.group(1)] = text[start:end].strip("\n")
    progress = sections.get("本周进展和问题", "").strip("\n")
    image_paths: list[str] = []
    progress_lines: list[str] = []
    in_images = False
    for line in progress.splitlines():
        stripped = line.strip()
        if re.fullmatch(r"图片[:：]\s*", stripped):
            in_images = True
            continue
        if in_images:
            bullet = re.match(r"^[-*]\s+(.+)$", stripped)
            if bullet:
                raw = bullet.group(1).strip()
                path = Path(raw).expanduser()
                if not path.is_absolute():
                    path = (input_path.parent / path).resolve()
                image_paths.append(str(path))
                continue
            if stripped.startswith("#"):
                continue
            if not stripped:
                continue
            in_images = False
        progress_lines.append(line)
    return {
        "goals": sections.get("本周目标和方案", "").strip("\n"),
        "progress": "\n".join(progress_lines).strip("\n"),
        "images": image_paths,
        "next_plan": sections.get("下周研究计划", "").strip("\n"),
        "other": sections.get("其他", "无").strip("\n") or "无",
    }


def json_dump(data: Any) -> None:
    print(json.dumps(data, ensure_ascii=False, indent=2))


def set_run_font(run, *, bold: bool = False, size=BODY_SIZE) -> None:
    run.font.name = BODY_FONT
    r_fonts = get_or_add_rfonts(run)
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = size
    run.bold = bold


def get_or_add_rfonts(run):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    return r_fonts


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._p.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    size=BODY_SIZE,
    alignment: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    clear_paragraph(paragraph)
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, size=size)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    *,
    bold: bool = False,
    size=BODY_SIZE,
    alignment: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.LEFT,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        set_paragraph_text(new_para, text, bold=bold, size=size, alignment=alignment)
    else:
        new_para.alignment = alignment
    return new_para


def clear_between(start: Paragraph, end: Paragraph) -> None:
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node


def find_heading(doc: Document, label: str) -> Paragraph:
    for para in doc.paragraphs:
        if normalize_space(para.text) == label:
            return para
    raise SystemExit(f"模板中未找到标题：{label}")


def replace_section(
    doc: Document,
    heading: str,
    next_heading: str,
    content: str,
    *,
    images: list[str] | None = None,
) -> None:
    start = find_heading(doc, heading)
    end = find_heading(doc, next_heading)
    clear_between(start, end)
    cursor = start
    lines = content.splitlines() if content else []
    if not lines and not images:
        lines = [""]
    for line in lines:
        cursor = insert_paragraph_after(cursor, line, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for image_path in images or []:
        path = Path(image_path)
        if not path.exists():
            raise SystemExit(f"图片不存在：{path}")
        image_para = insert_paragraph_after(cursor, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        image_para.add_run().add_picture(str(path), width=Emu(usable_width))
        cursor = image_para


def remove_extra_cell_paragraphs(cell) -> None:
    for para in list(cell.paragraphs)[1:]:
        para._element.getparent().remove(para._element)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    remove_extra_cell_paragraphs(cell)
    first = cell.paragraphs[0]
    clear_paragraph(first)
    parts = text.split("\n")
    paragraphs = [first]
    for _ in parts[1:]:
        paragraphs.append(cell.add_paragraph())
    for para, part in zip(paragraphs, parts):
        para.alignment = alignment
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(part)
        set_run_font(run, bold=bold, size=BODY_SIZE)


def draft_literature_from_paper(paper: dict[str, Any]) -> dict[str, str]:
    title = paper.get("title", "该论文")
    citation = paper.get("citation") or f"{title}. DOI: {paper.get('doi', '')}."
    venue = paper.get("venue", "")
    return {
        "citation": citation,
        "corresponding_author_unit": "需根据论文全文或出版社页面进一步确认通讯作者单位；当前元数据来源为 "
        + (venue or paper.get("source", "Crossref"))
        + "。",
        "research_purpose": f"论文围绕“{title}”展开，目标是提升相关红外光谱传感或计算光谱系统在小型化、谱信息获取和算法反演方面的能力。该方向与当前 16 通道红外气体传感工作具有方法上的相似性：前端通道提供有限维响应，后端需要通过标定、校正和算法提取有效特征。",
        "research_methods": "作者基于红外光谱测量或计算光谱框架建立系统模型，并结合器件响应、光谱编码、数据处理和算法重建/识别流程进行验证。阅读时应重点关注其响应矩阵标定、噪声处理、训练数据构建、评价指标和与传统光谱仪或基准算法的对比方式。",
        "conclusions": "论文说明，红外光谱传感系统的性能不仅由硬件通道数量决定，也受通道响应差异性、标定精度、噪声水平和后端算法约束影响。该结论可为有限通道红外传感系统的设计和数据处理提供参考。",
        "lessons": "对当前课题的启发在于：16 通道输出可以看作对气体吸收谱的压缩编码，后续应把暗噪声扣除、响应矩阵校正、归一化、特征提取和 PCA/LDA 或重建算法放在统一流程中评估，而不是只比较原始电压曲线。",
        "critique": "该论文的实验条件、通道结构和样本分布可能与当前固定 16 通道电路板不同，因此不能直接套用其误差水平或模型参数。后续需要重点验证通道稳定性、串扰、背景吸收和训练集覆盖度，并确认重建或分类算法是否真正提升气体识别准确率。",
    }


def replace_literature_table(doc: Document, literature: dict[str, str]) -> None:
    if not doc.tables:
        raise SystemExit("模板中没有文献精读表格。")
    table = doc.tables[0]
    if len(table.rows) < 7:
        raise SystemExit("文献精读表格行数不足。")
    set_cell_text(
        table.rows[0].cells[0],
        "参考文献基本信息【若尚未取得卷号和页码，给出DOI】",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(table.rows[0].cells[2], "通讯作者单位", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[1].cells[0], literature["citation"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(
        table.rows[1].cells[2],
        literature["corresponding_author_unit"],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    rows = [
        ("研究目的", "research_purpose"),
        ("研究方法", "research_methods"),
        ("论文结论", "conclusions"),
        ("从论文中学习到的内容", "lessons"),
        ("批判性\n思维", "critique"),
    ]
    for offset, (label, key) in enumerate(rows, start=2):
        set_cell_text(table.rows[offset].cells[0], label, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[offset].cells[1], literature[key], alignment=WD_ALIGN_PARAGRAPH.LEFT)


def parse_week_end_from_doc(path: Path, report_prefix: str | None = None) -> dt.date | None:
    try:
        doc = Document(path)
    except Exception:
        return None
    week = extract_week_text(doc)
    match = WEEK_RE.search(week)
    if match:
        start_year, _sm, _sd, end_year, em, ed = match.groups()
        year = int(end_year or start_year)
        try:
            return dt.date(year, int(em), int(ed))
        except ValueError:
            pass
    parsed = parse_report_date(path, report_prefix)
    return parsed[0] if parsed else None


def week_strings(
    report_dir: Path, week_ending: str | None, report_prefix: str
) -> tuple[dt.date, dt.date, str, str]:
    if week_ending:
        end = dt.date.fromisoformat(week_ending)
        start = end - dt.timedelta(days=7)
    else:
        latest = latest_report(report_dir, report_prefix)
        last_end = parse_week_end_from_doc(latest, report_prefix)
        if not last_end:
            raise SystemExit(f"无法从最新周报推断周次：{latest}")
        start = last_end
        end = last_end + dt.timedelta(days=7)
    if start.year == end.year:
        week_text = f"周次：{start.year}.{start.month}.{start.day}-{end.month}.{end.day}"
    else:
        week_text = (
            f"周次：{start.year}.{start.month}.{start.day}-{end.year}.{end.month}.{end.day}"
        )
    filename_date = f"{end.year % 100:02d}{end.month:02d}{end.day:02d}"
    return start, end, week_text, filename_date


def unique_output_path(output_dir: Path, filename_date: str, report_prefix: str) -> Path:
    base = output_dir / f"{report_prefix}-{filename_date}.docx"
    if not base.exists():
        return base
    version = 2
    while True:
        candidate = output_dir / f"{report_prefix}-{filename_date}-v{version}.docx"
        if not candidate.exists():
            return candidate
        version += 1


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def literature_from_args(args, paper: dict[str, Any]) -> dict[str, str]:
    if args.literature:
        literature = load_json(Path(args.literature).expanduser())
    else:
        literature = draft_literature_from_paper(paper)
    required = [
        "citation",
        "corresponding_author_unit",
        "research_purpose",
        "research_methods",
        "conclusions",
        "lessons",
        "critique",
    ]
    missing = [key for key in required if not normalize_space(literature.get(key, ""))]
    if missing:
        raise SystemExit("文献精读 JSON 缺少字段：" + ", ".join(missing))
    return {key: normalize_space(literature[key]) for key in required}


def find_render_docx() -> Path | None:
    explicit = os.environ.get("WEEKLY_REPORT_RENDER_DOCX")
    if explicit:
        path = Path(explicit).expanduser()
        if path.exists():
            return path
    patterns = [
        DEFAULT_RENDER_GLOB,
    ]
    matches: list[Path] = []
    for pattern in patterns:
        if "*" in str(pattern):
            matches.extend(sorted(Path.home().glob(str(pattern.relative_to(Path.home())))))
        elif pattern.exists():
            matches.append(pattern)
    return matches[-1] if matches else None


def render_docx(docx_path: Path) -> dict[str, Any]:
    renderer = find_render_docx()
    if not renderer:
        return {"status": "skipped", "reason": "未找到 render_docx.py"}
    out_dir = Path(tempfile.mkdtemp(prefix="weekly-report-render-"))
    env = os.environ.copy()
    env.setdefault("TMPDIR", "/private/tmp")
    cmd = [
        sys.executable,
        str(renderer),
        str(docx_path),
        "--output_dir",
        str(out_dir),
        "--emit_pdf",
    ]
    try:
        result = subprocess.run(cmd, text=True, capture_output=True, timeout=120, env=env)
    except Exception as exc:
        return {"status": "failed", "reason": str(exc), "output_dir": str(out_dir)}
    pages = sorted(str(path) for path in out_dir.glob("page*.png"))
    return {
        "status": "ok" if result.returncode == 0 and pages else "failed",
        "output_dir": str(out_dir),
        "pages": pages,
        "stdout": result.stdout.strip(),
        "stderr": result.stderr.strip(),
        "returncode": result.returncode,
    }


def generate_report(args) -> dict[str, Any]:
    report_dir = report_dir_from_args(args.report_dir)
    report_prefix = report_prefix_from_args(args.report_prefix)
    used_file = used_file_for(report_dir, args.used_file)
    output_dir = Path(args.output_dir).expanduser() if args.output_dir else report_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    template = Path(args.template).expanduser() if args.template else latest_report(report_dir, report_prefix)
    input_path = input_path_from_args(args, report_dir)
    input_data = parse_weekly_input(input_path)
    paper: dict[str, Any]
    if args.paper:
        paper = load_json(Path(args.paper).expanduser())
    elif args.auto_paper:
        paper = select_paper(report_dir, used_file, report_prefix)
    else:
        paper = {}
    literature = literature_from_args(args, paper)
    _start, _end, week_text, filename_date = week_strings(
        report_dir, args.week_ending, report_prefix
    )
    output_path = unique_output_path(output_dir, filename_date, report_prefix)

    shutil.copy2(template, output_path)
    doc = Document(output_path)

    for para in doc.paragraphs:
        if normalize_space(para.text).startswith("周次"):
            set_paragraph_text(para, week_text, bold=True, size=HEADER_SIZE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            break
    else:
        raise SystemExit("模板中未找到周次段落。")

    replace_section(
        doc,
        "一、本周目标和方案",
        "二、本周进展和问题",
        input_data["goals"],
    )
    replace_section(
        doc,
        "二、本周进展和问题",
        "三、下周研究计划",
        input_data["progress"],
        images=input_data["images"],
    )
    replace_section(
        doc,
        "三、下周研究计划",
        "四、其他",
        input_data["next_plan"],
    )
    replace_section(
        doc,
        "四、其他",
        "本周文献精读",
        input_data["other"],
    )
    replace_literature_table(doc, literature)
    doc.save(output_path)

    used_synced = False
    if output_dir.resolve() == report_dir.resolve():
        sync_used(report_dir, used_file, report_prefix)
        used_synced = True

    render = {"status": "skipped", "reason": "--no-render"}
    if not args.no_render:
        render = render_docx(output_path)
    return {
        "output": str(output_path),
        "template": str(template),
        "input": str(input_path),
        "report_prefix": report_prefix,
        "week": week_text,
        "used_file": str(used_file),
        "used_synced": used_synced,
        "render": render,
    }


def cmd_sync_used(args) -> None:
    report_dir = report_dir_from_args(args.report_dir)
    report_prefix = report_prefix_from_args(args.report_prefix)
    used_file = used_file_for(report_dir, args.used_file)
    records = sync_used(report_dir, used_file, report_prefix)
    json_dump({"used_file": str(used_file), "report_prefix": report_prefix, "records": len(records)})


def cmd_select_paper(args) -> None:
    report_dir = report_dir_from_args(args.report_dir)
    report_prefix = report_prefix_from_args(args.report_prefix)
    used_file = used_file_for(report_dir, args.used_file)
    paper = select_paper(report_dir, used_file, report_prefix)
    if args.out:
        out = Path(args.out).expanduser()
        out.write_text(json.dumps(paper, ensure_ascii=False, indent=2), encoding="utf-8")
    json_dump(paper)


def cmd_parse_input(args) -> None:
    report_dir = report_dir_from_args(args.report_dir)
    report_prefix = report_prefix_from_args(args.report_prefix)
    input_path = input_path_from_args(args, report_dir)
    data = parse_weekly_input(input_path)
    data["input"] = str(input_path)
    data["report_prefix"] = report_prefix
    json_dump(data)


def cmd_generate(args) -> None:
    json_dump(generate_report(args))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="生成通用周报 DOCX 并维护 used.md",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent(
            """\
            示例：
              weekly_report.py sync-used
              weekly_report.py select-paper --out /tmp/weekly-paper.json
              weekly_report.py parse-input
              weekly_report.py generate --auto-paper
              weekly_report.py generate --report-prefix student-weekly-report --auto-paper
              weekly_report.py generate --input weekly.md --paper paper.json --literature literature.json
            """
        ),
    )
    sub = parser.add_subparsers(dest="command", required=True)

    def add_common(p):
        p.add_argument("--report-dir", help="周报目录，默认读取 WEEKLY_REPORT_DIR 或当前目录")
        p.add_argument("--used-file", help="used.md 路径，默认是周报目录下 used.md")
        p.add_argument(
            "--report-prefix",
            help="周报文件名前缀，默认读取 WEEKLY_REPORT_PREFIX 或 weekly-report",
        )

    p = sub.add_parser("sync-used", help="扫描周报目录并重写 used.md")
    add_common(p)
    p.set_defaults(func=cmd_sync_used)

    p = sub.add_parser("select-paper", help="从 Crossref 自动选择未重复相关文献")
    add_common(p)
    p.add_argument("--out", help="把选中文献 JSON 写入指定路径")
    p.set_defaults(func=cmd_select_paper)

    p = sub.add_parser("parse-input", help="解析每周输入 Markdown")
    add_common(p)
    p.add_argument(
        "--input",
        help="每周输入 Markdown 文件；默认读取周报目录下 weekly-input.md",
    )
    p.set_defaults(func=cmd_parse_input)

    p = sub.add_parser("generate", help="生成新的周报 DOCX")
    add_common(p)
    p.add_argument(
        "--input",
        help="每周输入 Markdown 文件；默认读取周报目录下 weekly-input.md",
    )
    p.add_argument("--paper", help="select-paper 输出的文献 JSON")
    p.add_argument("--literature", help="文献精读内容 JSON")
    p.add_argument("--auto-paper", action="store_true", help="未提供 --paper 时自动选择文献")
    p.add_argument("--template", help="指定模板 DOCX，默认使用最新周报")
    p.add_argument("--output-dir", help="输出目录，默认写入周报目录")
    p.add_argument("--week-ending", help="本周结束日期 YYYY-MM-DD，默认从最新周报往后推一周")
    p.add_argument("--no-render", action="store_true", help="跳过 DOCX 渲染检查")
    p.set_defaults(func=cmd_generate)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    args.func(args)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
